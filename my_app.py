from operator import length_hint
from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

document.add_picture(
    'me.jpg',
    width=Inches(2.0), 
    height=Inches(2.0)
) 

welcome = print('Welcome to the CV builder. Follow these steps to complete your CV. Press Control + C to quit.')
name = input('What is your full name? ')
speak('Hello ' + name + ' How are you today? ')
speak('What is your phone number? ')
phone_number = input('What is your phone number? ')
speak(phone_number)
email = input('What is your email address ? ')
speak(email)

document.add_paragraph(
    name + ' | ' + phone_number + ' | '  + email)

document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? '))

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Please enter a company name ')
start_date = input ('What date did you start working for this company? ')
end_date = input ('When did your employement end? ')
experience = input ('Tell us about your experience at ' + company + '? ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True
p.add_run(experience)

# additional work experience 
while True:
    has_more_experience = input('Do you have more work experience? yes or no ')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Please enter the company name ')
        start_date = input ('What date did you start working for this company? ')
        end_date = input ('When did your employement end? ')
        experience = input ('Tell us about your experience at ' + company + '? ')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + '-' + end_date + '\n').italic = True
        p.add_run(experience)
    else:
        end = input('Great. Lets move onto the skills section. Press enter to continue')
        print(end)
        break

# Skills
document.add_heading('Skills')
skill = input(name + ', lets add your skills. Please input a skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have another skill? yes or no ')
    if has_more_skills.lower() == 'yes':
        skill = input('Please enter another skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        end = input('Congratulations! Your CV is now complete. Press Enter to finish')
        print(end)
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV created using Python"


document.save('cv.docx')
